from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import io
import locale

try:
    locale.setlocale(locale.LC_ALL, "es_CL.UTF-8")
except Exception:
    pass


def _fmt(n: int) -> str:
    return f"${n:,.0f}".replace(",", ".")


def _numero_a_palabras(n: int) -> str:
    """Convierte un entero en texto para la cláusula TERCERO."""
    # Implementación básica para rangos de 1M a 10M
    millones = n // 1_000_000
    resto    = n % 1_000_000
    miles    = resto // 1_000
    unidades = resto % 1_000

    partes = []
    if millones == 1:
        partes.append("un millón")
    elif millones > 1:
        partes.append(f"{millones} millones")
    if miles > 0:
        partes.append(f"{miles} mil")
    if unidades > 0:
        partes.append(str(unidades))
    return " ".join(partes) + " pesos"


def _add_para(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _add_mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    """parts = [(texto, bold)]"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    for text, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


def _fecha_larga(d: date) -> str:
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    return f"{d.day:02d} de {meses[d.month - 1]} de {d.year}"


def generar_finiquito_docx(
    empleado: dict,
    finiquito: dict,
    cuotas: list[dict],
    ciudad_notaria: str,
    fecha_firma: date | None = None,
) -> bytes:
    if fecha_firma is None:
        fecha_firma = date.today()

    doc = Document()

    # Márgenes 2.5 cm
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(85)
        section.right_margin  = Pt(85)

    # Título
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(14)
    r = t.add_run("Finiquito de Contrato de Trabajo")
    r.bold = True
    r.font.size = Pt(14)

    emp   = empleado
    fin   = finiquito
    anos  = fin["anos_servicio"]
    fer   = fin["feriado_proporcional"]
    total = fin["total"]
    base  = fin["base_indemnizacion"]

    n_cuotas = len(cuotas)
    cuenta   = emp.get("cuenta_bancaria", {})

    # ENCABEZADO
    causal_texto = {
        "161_1": "Necesidades de la empresa, causal señalada en el artículo 161, inciso 1, del Código del Trabajo",
        "161_2": "Desahucio del empleador, causal señalada en el artículo 161, inciso 2, del Código del Trabajo",
        "159_1": "Mutuo acuerdo de las partes, causal señalada en el artículo 159 N°1 del Código del Trabajo",
        "159_2": "Renuncia voluntaria, causal señalada en el artículo 159 N°2 del Código del Trabajo",
        "159_4": "Vencimiento del plazo convenido, causal señalada en el artículo 159 N°4 del Código del Trabajo",
        "159_5": "Conclusión del trabajo o servicio que dio origen al contrato, causal señalada en el artículo 159 N°5 del Código del Trabajo",
    }.get(fin["causal"], fin["causal"])

    _add_mixed(doc, [
        (f"En {ciudad_notaria}, a {_fecha_larga(fecha_firma)}, entre ", False),
        ("AGENCIA CONVIERTE SPA", True),
        (f", RUT 77.450.452-4, representada legalmente por don ", False),
        ("KARIM EMILIO ABUHADBA PÉREZ", True),
        (", RUT 16.608.090-8, domiciliado en Del Pucará 50, oficina 511-B, comuna de Machalí, en adelante también «el empleador», por una parte; y, por la otra, doña ", False),
        (emp["nombre"].upper(), True),
        (f", RUT {emp['rut']}, en adelante también «la trabajadora», se deja testimonio y se ha acordado el finiquito que consta de las siguientes cláusulas:", False),
    ])

    # PRIMERO
    fecha_inicio  = fin["fecha_inicio"]
    fecha_termino = fin["fecha_termino"]
    _add_mixed(doc, [
        ("PRIMERO: ", True),
        (f"La trabajadora prestó servicios al empleador desde el {fecha_inicio} hasta el {fecha_termino}, fecha esta última en que su contrato de trabajo ha terminado por ", False),
        (causal_texto, True),
        (".", False),
    ])

    # SEGUNDO
    _add_mixed(doc, [
        ("SEGUNDO: ", True),
        (f"Doña {emp['nombre'].upper()} declara recibir de parte de AGENCIA CONVIERTE SPA la suma de {_fmt(total)}, según la liquidación que se señala a continuación:", False),
    ])

    # Tabla de haberes
    if fer.get("dias_pendientes") is not None:
        fer_label = f"Feriado proporcional ({fer['dias_pendientes']:.2f} días hábiles)"
    else:
        dias_tomados = fer.get("dias_tomados", "")
        fer_label = f"Feriado proporcional (según liquidación)"

    items = [
        (fer_label, fer["monto"]),
    ]
    if fin["indemnizacion_anos"] > 0:
        label_anos = f"Indemnización por años de servicio ({anos['anos_indemnizacion']} {'año' if anos['anos_indemnizacion'] == 1 else 'años'})"
        items.append((label_anos, fin["indemnizacion_anos"]))
    if fin["indemnizacion_aviso"] > 0:
        items.append(("Indemnización sustitutiva del aviso previo", fin["indemnizacion_aviso"]))

    for label, monto in items:
        _add_para(doc, f"    -  {label} {'.' * max(1, 55 - len(label))} {_fmt(monto)}")

    _add_para(doc, f"    TOTAL A CANCELAR {'.' * 44} {_fmt(total)}", bold=True)

    # Cuotas
    cuotas_texto = "en 1 cuota" if n_cuotas == 1 else f"en {n_cuotas} cuotas"
    _add_mixed(doc, [
        (f"Las partes acuerdan expresamente que la suma total indicada precedentemente, ascendente a {_fmt(total)}, será pagada por el empleador {cuotas_texto}, de la siguiente forma:", False),
    ])

    for c in cuotas:
        _add_para(doc, f"    -  Cuota {c['numero']}, {c['fecha']}: {_fmt(c['monto'])}")

    banco   = cuenta.get("banco", "___________")
    tipo    = cuenta.get("tipo", "Cuenta RUT")
    numero  = cuenta.get("numero", "___________")
    titular = emp["nombre"]
    rut_emp = emp["rut"]

    _add_para(doc, f"El pago de cada cuota se efectuará mediante transferencia bancaria a la siguiente cuenta: {banco}, {tipo} N° {numero}, a nombre de {titular}, RUT {rut_emp}.")

    _add_para(doc, "Las partes dejan expresa constancia de que el presente finiquito constituye título ejecutivo de conformidad con el artículo 462 del Código del Trabajo, y que el no pago oportuno de cualquiera de las cuotas pactadas hará exigible de inmediato la totalidad del saldo insoluto, sin necesidad de requerimiento previo.")

    _add_para(doc, f"Doña {emp['nombre'].upper()} declara haber analizado y estudiado detenidamente dicha liquidación, aceptándola en todas sus partes, sin tener observación alguna que formularle.")

    # TERCERO
    total_en_palabras = _numero_a_palabras(total)
    _add_mixed(doc, [
        ("TERCERO: ", True),
        (f"En consecuencia, el empleador paga a doña {emp['nombre'].upper()} la suma de {_fmt(total)} ({total_en_palabras}), de conformidad con el calendario de cuotas establecido anteriormente. Las partes dejan constancia de que la referida suma cubre el total de los haberes especificados en la liquidación señalada en el numerando SEGUNDO del presente finiquito.", False),
    ])

    # CUARTO - cotizaciones
    _add_mixed(doc, [
        ("CUARTO: ", True),
        ("El empleador declara y acredita, mediante los certificados respectivos que en este acto se exhiben, encontrarse al día en el pago íntegro de las cotizaciones previsionales, de salud y del seguro de cesantía de la trabajadora, hasta el último día del mes anterior al término de la relación laboral, dando así cumplimiento a lo dispuesto en el artículo 162, incisos quinto y siguientes, del Código del Trabajo.", False),
    ])

    # QUINTO
    _add_mixed(doc, [
        ("QUINTO: ", True),
        (f"Doña {emp['nombre'].upper()} deja constancia que durante el tiempo que prestó servicios a AGENCIA CONVIERTE SPA recibió oportunamente el total de las remuneraciones, beneficios y demás prestaciones convenidas de acuerdo a su contrato de trabajo, clase de trabajo ejecutado y disposiciones legales pertinentes, y que en tal virtud el empleador nada le adeuda por tales conceptos, ni por horas extraordinarias, asignación familiar, feriado, indemnización por años de servicios, imposiciones previsionales, así como por ningún otro concepto, ya sea legal o contractual, derivado de la prestación de sus servicios. Declara, asimismo, que no tiene reclamo alguno que formular en contra de AGENCIA CONVIERTE SPA, renunciando a todas las acciones que pudieran emanar del contrato que los vinculó.", False),
    ])

    # SEXTO
    _add_mixed(doc, [
        ("SEXTO: ", True),
        (f"En virtud de lo anteriormente expuesto, doña {emp['nombre'].upper()} manifiesta expresamente que AGENCIA CONVIERTE SPA nada le adeuda en relación con los servicios prestados, con el contrato de trabajo o con motivo de la terminación del mismo, por lo que libre y espontáneamente, y con el pleno y cabal conocimiento de sus derechos, otorga a su empleador el más amplio, completo, total y definitivo finiquito por los servicios prestados o la terminación de ellos, ya diga relación con remuneraciones, cotizaciones previsionales, de seguridad social o de salud, subsidios, beneficios contractuales adicionales a las remuneraciones, indemnizaciones, compensaciones, o con cualquiera causa o concepto.", False),
    ])

    # SÉPTIMO
    _add_mixed(doc, [
        ("SÉPTIMO: ", True),
        (f"Asimismo, declara la trabajadora que, en todo caso y a todo evento, renuncia expresamente a cualquier derecho, acción o reclamo que eventualmente tuviere o pudiere corresponderle en contra del empleador, en relación directa o indirecta con su contrato de trabajo, con los servicios prestados, con la terminación del referido contrato o de dichos servicios.", False),
    ])

    # OCTAVO - Ley 21.389 alimentos
    _add_mixed(doc, [
        ("OCTAVO: ", True),
        ("Según la Ley 21.389, el empleador declara bajo juramento que la trabajadora no se encuentra sujeta a retención judicial por pensión de alimentos, según se acredita con la exhibición de las 3 últimas liquidaciones de sueldo, liberando en este acto al ministro de fe autorizante de cualquier responsabilidad por el no pago de ésta.", False),
    ])

    _add_para(doc, "Para constancia, las partes firman el presente finiquito en tres ejemplares, quedando uno en poder de cada una de ellas y el tercero en poder del ministro de fe. En cumplimiento de la legislación vigente, la trabajadora lo lee, firma y lo ratifica ante el ministro de fe que lo autoriza.")

    # FIRMAS
    doc.add_paragraph()
    doc.add_paragraph()

    tabla_firmas = doc.add_table(rows=1, cols=2)
    tabla_firmas.style = "Table Grid"
    celdas = tabla_firmas.rows[0].cells
    for celda in celdas:
        celda._element.get_or_add_tcPr()

    def firma_celda(celda, lineas):
        p = celda.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(50)
        for i, linea in enumerate(lineas):
            r = p.add_run(("\n" if i > 0 else "") + linea)
            r.bold = i == 0
            r.font.size = Pt(10)

    firma_celda(celdas[0], [
        "_" * 35,
        emp["nombre"].upper(),
        f"RUT {emp['rut']}",
        "Trabajadora",
    ])
    firma_celda(celdas[1], [
        "_" * 35,
        "KARIM EMILIO ABUHADBA PÉREZ",
        "RUT 16.608.090-8",
        "en representación de AGENCIA CONVIERTE SPA",
        "RUT 77.450.452-4",
        "Empleador",
    ])

    doc.add_paragraph()
    p_fe = doc.add_paragraph()
    p_fe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fe.paragraph_format.space_before = Pt(50)
    r = p_fe.add_run("_" * 50 + "\nMinistro de fe\nNombre: _______________________________\nCalidad (Notario / Inspector del Trabajo): _______________________________")
    r.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
